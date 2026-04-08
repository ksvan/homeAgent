"""
Tools MCP server — bash, python exec, web scrape, web search.

Tools are registered conditionally based on feature flags so the server
always starts cleanly regardless of which features are enabled.
"""
from __future__ import annotations

import ipaddress
import logging
import re
import shutil
import socket
import uuid
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import quote, urlparse

from fastmcp import FastMCP

from app.config import Settings, get_settings
from app.shell import ALWAYS_BLOCKED, DEFAULT_ALLOWED, run_command

logger = logging.getLogger(__name__)


def _is_ssrf_blocked(url: str) -> bool:
    """Return True if the URL's hostname resolves to a private/reserved address.

    Blocks requests to loopback, private RFC-1918, link-local, and multicast
    ranges to prevent SSRF attacks against internal services.
    Fails closed (returns True) on resolution errors.
    """
    try:
        parsed = urlparse(url)
        host = parsed.hostname or ""
        addrs = socket.getaddrinfo(host, None)
        for *_, sockaddr in addrs:
            ip = ipaddress.ip_address(sockaddr[0])
            if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_multicast:
                return True
    except Exception:
        return True  # fail-closed on resolution error
    return False


mcp = FastMCP(
    name="tools",
    instructions=(
        "Sandboxed execution tools: bash shell commands, Python scripts, "
        "web page scraping, and web search. "
        "All code execution runs in an isolated workspace directory."
    ),
)

_TMP_DIR = ".agent_tmp"
_CLEANUP_AGE_HOURS = 24


def _cleanup_old_runs(workspace: Path) -> None:
    tmp_base = workspace / _TMP_DIR
    if not tmp_base.exists():
        return
    cutoff = datetime.now(timezone.utc).timestamp() - _CLEANUP_AGE_HOURS * 3600
    for d in tmp_base.iterdir():
        if d.is_dir() and d.stat().st_mtime < cutoff:
            shutil.rmtree(d, ignore_errors=True)


def register_tools(settings: Settings) -> None:
    """Register enabled tools on the FastMCP server based on feature flags."""

    workspace = Path(settings.workspace_dir)
    passthrough_env = settings.passthrough_env_dict()

    # ------------------------------------------------------------------
    # Bash tool
    # ------------------------------------------------------------------
    if settings.feature_bash:
        _cmd_list = settings.bash_allowed_commands_list()
        if _cmd_list:
            allowed = frozenset(_cmd_list) - ALWAYS_BLOCKED
        else:
            allowed = DEFAULT_ALLOWED

        bash_max_timeout = settings.bash_max_timeout_seconds
        bash_max_output = settings.bash_max_output_bytes

        @mcp.tool
        async def run_bash_command(
            argv: list[str],
            cwd: str = ".",
            timeout_s: int = 30,
        ) -> str:
            """Run an allowlisted command in the household workspace directory.

            Only plain argument lists are accepted — no shell metacharacters, pipes,
            redirects, or variable expansion (shell=False).

            IMPORTANT — confirm before write operations:
            - READ-ONLY (ls, cat, grep, find, git status, head, etc.) → run immediately.
            - WRITE / MODIFY (cp, mv, touch, mkdir, git commit, etc.) → ask for confirmation first.

            Args:
                argv: Command as a list of strings, e.g. ["grep", "-r", "error", "logs/"].
                      Never include shell operators (|, >, <, &&, ;).
                cwd:  Working directory relative to the workspace root. Must not escape
                      the workspace (no '..' traversal).
                timeout_s: Seconds before the process is killed. Max 300, default 30.
            """
            t = min(timeout_s, bash_max_timeout)
            workspace.mkdir(parents=True, exist_ok=True)
            result = await run_command(
                argv=argv,
                cwd=cwd,
                timeout_s=t,
                workspace_dir=workspace,
                allowed_commands=allowed,
                max_output_bytes=bash_max_output,
                extra_env=passthrough_env or None,
            )
            parts: list[str] = []
            if result.stdout:
                parts.append(result.stdout.rstrip())
            if result.stderr:
                parts.append(f"[stderr]\n{result.stderr.rstrip()}")
            if result.truncated:
                parts.append("[Output truncated to size limit]")
            if not result.ok:
                parts.append(f"[exit {result.exit_code}]")
            return "\n".join(parts) if parts else "(no output)"

        logger.info("Tools MCP: bash tool registered")

    # ------------------------------------------------------------------
    # Python exec tool
    # ------------------------------------------------------------------
    if settings.feature_python:
        py_max_timeout = settings.python_max_timeout_seconds
        py_max_output = settings.python_max_output_bytes

        @mcp.tool
        async def run_python_script(
            code: str,
            files: dict[str, str] | None = None,
            timeout_s: int = 30,
        ) -> str:
            """Write and execute a Python script in an isolated workspace directory.

            Use this when you need to:
            - Perform calculations or data transformations
            - Process files from the workspace
            - Generate output files (charts, reports, processed data)
            - Run multi-step logic that is cleaner as a script than a shell command

            The script runs with read/write access to the workspace only.

            IMPORTANT — confirm before running scripts that write files.

            Args:
                code:    The Python script to run (written as main.py).
                files:   Optional helper files as {filename: content}.
                timeout_s: Seconds before the process is killed. Max 300, default 30.
            """
            t = min(timeout_s, py_max_timeout)
            workspace.mkdir(parents=True, exist_ok=True)
            _cleanup_old_runs(workspace)

            run_id = uuid.uuid4().hex[:12]
            run_dir = workspace / _TMP_DIR / run_id
            run_dir.mkdir(parents=True)

            (run_dir / "main.py").write_text(code, encoding="utf-8")
            input_names = {"main.py"}
            for name, content in (files or {}).items():
                safe_name = Path(name).name
                (run_dir / safe_name).write_text(content, encoding="utf-8")
                input_names.add(safe_name)

            result = await run_command(
                argv=["python3", "main.py"],
                cwd=str(run_dir.relative_to(workspace)),
                timeout_s=t,
                workspace_dir=workspace,
                allowed_commands=frozenset({"python3"}),
                max_output_bytes=py_max_output,
                extra_env=passthrough_env or None,
            )

            artifacts: list[str] = []
            for f in sorted(run_dir.rglob("*")):
                if f.is_file() and f.name not in input_names:
                    rel = f.relative_to(run_dir)
                    size = f.stat().st_size
                    artifacts.append(f"  {rel} ({size} bytes)")

            parts: list[str] = []
            if result.stdout:
                parts.append(result.stdout.rstrip())
            if result.stderr:
                parts.append(f"[stderr]\n{result.stderr.rstrip()}")
            if result.truncated:
                parts.append("[Output truncated to size limit]")
            if not result.ok:
                parts.append(f"[exit {result.exit_code}]")
            if artifacts:
                parts.append("Output files:\n" + "\n".join(artifacts))
            return "\n".join(parts) if parts else "(no output)"

        logger.info("Tools MCP: python tool registered")

    # ------------------------------------------------------------------
    # Web scrape tool
    # ------------------------------------------------------------------
    if settings.feature_scrape:
        scrape_timeout = settings.scrape_timeout_seconds
        scrape_max = settings.scrape_max_content_bytes

        @mcp.tool
        async def scrape_web_page(
            url: str,
            timeout_s: int = 20,
        ) -> str:
            """Fetch a web page and return its readable text content.

            Use this when the user asks to look something up online, check a website,
            read an article, or get information from a specific URL. Only use this
            when the user explicitly asks for a web fetch.

            Returns the page title and main text content with boilerplate removed.

            Args:
                url:       The full URL to fetch (http or https).
                timeout_s: Request timeout in seconds (default 20, max 60).
            """
            import httpx
            from bs4 import BeautifulSoup

            t = min(timeout_s, scrape_timeout)

            try:
                parsed = urlparse(url)
                if parsed.scheme not in ("http", "https") or not parsed.netloc:
                    return f"Invalid URL: {url!r}. Only http/https URLs are supported."
            except Exception:
                return f"Could not parse URL: {url!r}"

            if _is_ssrf_blocked(url):
                return f"URL not allowed: '{parsed.hostname}' resolves to a private or reserved address."

            logger.info("Scraping URL: %s", url)

            try:
                async with httpx.AsyncClient(
                    follow_redirects=False,
                    timeout=t,
                    headers={
                        "User-Agent": "Mozilla/5.0 (compatible; HomeAgent/1.0)",
                        "Accept-Language": "en,*;q=0.5",
                    },
                ) as client:
                    response = await client.get(url)
                    response.raise_for_status()
            except httpx.HTTPStatusError as exc:
                return f"HTTP {exc.response.status_code} from {url}"
            except httpx.TimeoutException:
                return f"Request timed out after {t}s: {url}"
            except Exception as exc:
                return f"Failed to fetch {url}: {exc}"

            content_type = response.headers.get("content-type", "")
            if "text/html" not in content_type and "application/xhtml" not in content_type:
                text = response.text[:scrape_max]
                return f"[{content_type}]\n{text}"

            soup = BeautifulSoup(response.text, "html.parser")

            title = ""
            if soup.title and soup.title.string:
                title = soup.title.string.strip()

            for tag in soup(
                ["script", "style", "head", "nav", "footer", "header", "aside",
                 "form", "button", "iframe", "noscript"]
            ):
                tag.decompose()

            text = soup.get_text(separator="\n", strip=True)
            text = re.sub(r" {2,}", " ", text)
            text = re.sub(r"\n{3,}", "\n\n", text)
            text = text.strip()

            if len(text.encode("utf-8")) > scrape_max:
                text = text.encode("utf-8")[:scrape_max].decode("utf-8", errors="ignore")
                text += "\n\n[Content truncated]"

            result = f"# {title}\n\n{text}" if title else text
            logger.info("Scraped %s — %d chars returned", url, len(result))
            return result

        logger.info("Tools MCP: scrape tool registered")

    # ------------------------------------------------------------------
    # Web search tool
    # ------------------------------------------------------------------
    if settings.feature_search:
        search_max = settings.search_max_results

        @mcp.tool
        async def search_web(
            query: str,
            max_results: int = 5,
        ) -> str:
            """Search the web and return a list of results with title, URL, and excerpt.

            Use this when the user asks for current information, news, prices, events,
            how-to guides, or anything that might have changed since training data.

            Do NOT use this for a specific URL — use scrape_web_page instead.

            Args:
                query:       The search query. Be specific; include year or context if relevant.
                max_results: Number of results to return (1–10, default 5).
            """
            if not settings.tavily_api_key:
                return "Search is not available: TAVILY_API_KEY is not configured."

            n = min(max(1, max_results), search_max)
            logger.info("Web search: %r (max_results=%d)", query, n)

            try:
                from tavily import AsyncTavilyClient  # type: ignore[import-untyped]

                client = AsyncTavilyClient(api_key=settings.tavily_api_key)
                response = await client.search(query, max_results=n)
                results = response.get("results", [])
            except Exception as exc:
                logger.exception("Search failed")
                return f"Search failed: {exc}"

            if not results:
                return "No results found."

            lines: list[str] = []
            for i, r in enumerate(results, 1):
                lines.append(f"{i}. {r.get('title', '')}")
                lines.append(f"   {r.get('url', '')}")
                snippet = r.get("content", "")
                if snippet:
                    snippet = snippet[:400].rstrip()
                    lines.append(f"   {snippet}")
                lines.append("")

            return "\n".join(lines).rstrip()

        logger.info("Tools MCP: search tool registered")

    # ------------------------------------------------------------------
    # SharePoint tools
    # ------------------------------------------------------------------
    if settings.feature_sharepoint:
        sp_timeout = settings.sharepoint_timeout_seconds
        sp_max_file = settings.sharepoint_max_file_bytes
        sp_max_content = settings.sharepoint_max_content_bytes

        @mcp.tool
        async def sharepoint_list_files(
            site_url: str,
            folder_path: str = "/Shared Documents",
        ) -> str:
            """List files and subfolders in a SharePoint document library.

            Requires the site to allow anonymous/guest access. Use this to discover
            what files are available before downloading them.

            The site_url should be the base site URL, e.g.:
              https://tenant.sharepoint.com/sites/MySite

            For "anyone with the link" sharing links (/:f:/g/… style), use
            sharepoint_download_file directly instead — it follows redirects to
            the real file URL.

            Returns a list of files with name, size, modified date, and server-relative
            URL (pass the full URL constructed as site_url + server_relative_url to
            sharepoint_download_file).

            Args:
                site_url:    Base SharePoint site URL (no trailing slash).
                folder_path: Server-relative path to the folder, e.g.
                             "/sites/MySite/Shared Documents/Reports".
                             Defaults to "/Shared Documents".
            """
            import httpx

            try:
                parsed = urlparse(site_url)
                if parsed.scheme not in ("http", "https") or not parsed.netloc:
                    return f"Invalid site_url: {site_url!r}. Only http/https URLs are supported."
            except Exception:
                return f"Could not parse site_url: {site_url!r}"

            if _is_ssrf_blocked(site_url):
                return f"URL not allowed: '{urlparse(site_url).hostname}' resolves to a private or reserved address."

            headers = {
                "Accept": "application/json;odata=verbose",
                "User-Agent": "Mozilla/5.0 (compatible; HomeAgent/1.0)",
            }

            base = site_url.rstrip("/")
            encoded_path = quote(folder_path, safe="/")
            files_url = f"{base}/_api/web/GetFolderByServerRelativeUrl('{encoded_path}')/Files"
            folders_url = f"{base}/_api/web/GetFolderByServerRelativeUrl('{encoded_path}')/Folders"

            logger.info("SharePoint list: site=%s folder=%s", base, folder_path)

            try:
                async with httpx.AsyncClient(
                    follow_redirects=True,
                    timeout=sp_timeout,
                    headers=headers,
                ) as client:
                    files_resp = await client.get(files_url)
                    files_resp.raise_for_status()
                    folders_resp = await client.get(folders_url)
                    folders_resp.raise_for_status()
            except httpx.HTTPStatusError as exc:
                return f"HTTP {exc.response.status_code} from SharePoint — the site may require authentication or the folder path is wrong."
            except httpx.TimeoutException:
                return f"Request timed out after {sp_timeout}s."
            except Exception as exc:
                return f"Failed to reach SharePoint: {exc}"

            try:
                files_data = files_resp.json()
                folders_data = folders_resp.json()
            except Exception:
                return "SharePoint returned an unexpected response (not JSON)."

            if "error" in files_data:
                msg = files_data["error"].get("message", {})
                return f"SharePoint error: {msg.get('value', files_data['error'])}"

            lines: list[str] = [f"Contents of {folder_path}:", ""]

            folder_results = (folders_data.get("d", {}).get("results") or [])
            for f in folder_results:
                name = f.get("Name", "")
                if name not in ("Forms",):  # skip SP internal folder
                    lines.append(f"[folder] {name}/")

            file_results = (files_data.get("d", {}).get("results") or [])
            for f in file_results:
                name = f.get("Name", "")
                size = f.get("Length", "?")
                modified = (f.get("TimeLastModified") or "")[:10]
                rel_url = f.get("ServerRelativeUrl", "")
                full_url = f"{parsed.scheme}://{parsed.netloc}{rel_url}"
                try:
                    size_kb = f"{int(size) // 1024} KB"
                except (ValueError, TypeError):
                    size_kb = str(size)
                lines.append(f"{name}  ({size_kb}, {modified})  {full_url}")

            if not folder_results and not file_results:
                lines.append("(empty folder or access denied)")

            return "\n".join(lines)

        @mcp.tool
        async def sharepoint_download_file(
            file_url: str,
        ) -> str:
            """Download a file from SharePoint and return its text content.

            Supports:
            - .docx — extracts paragraphs and tables as plain text
            - .txt, .csv, .md — returns raw text
            - .pdf — returns a note (PDF parsing not available in this service)
            - other binary types — returns file name and size only

            Also works with SharePoint "anyone with the link" sharing URLs
            (/:w:/g/… or /:f:/g/… style) — redirects are followed automatically
            to reach the real file.

            Requires the file or site to allow anonymous/guest access.

            Note: DOCX extraction covers paragraphs and tables only. Embedded
            images, charts, and SmartArt are not extracted.

            Args:
                file_url: Full HTTPS URL of the file to download.
            """
            import httpx

            try:
                parsed = urlparse(file_url)
                if parsed.scheme not in ("http", "https") or not parsed.netloc:
                    return f"Invalid URL: {file_url!r}. Only http/https URLs are supported."
            except Exception:
                return f"Could not parse URL: {file_url!r}"

            if _is_ssrf_blocked(file_url):
                return f"URL not allowed: '{parsed.hostname}' resolves to a private or reserved address."

            filename = Path(parsed.path).name
            ext = Path(parsed.path).suffix.lower()

            logger.info("SharePoint download: %s", file_url)

            try:
                async with httpx.AsyncClient(
                    follow_redirects=True,
                    timeout=sp_timeout,
                    headers={"User-Agent": "Mozilla/5.0 (compatible; HomeAgent/1.0)"},
                ) as client:
                    response = await client.get(file_url)
                    response.raise_for_status()
                    # Stream up to the size cap
                    raw = response.content
            except httpx.HTTPStatusError as exc:
                return f"HTTP {exc.response.status_code} fetching file — the file may require authentication."
            except httpx.TimeoutException:
                return f"Request timed out after {sp_timeout}s."
            except Exception as exc:
                return f"Failed to fetch file: {exc}"

            if len(raw) > sp_max_file:
                return f"File too large: {len(raw) // 1024} KB exceeds the {sp_max_file // 1_000_000} MB limit."

            # Re-check extension from final URL (after redirects) if filename changed
            final_path = Path(urlparse(str(response.url)).path).suffix.lower()
            if final_path:
                ext = final_path

            content_type = response.headers.get("content-type", "")

            if ext == ".docx" or "officedocument.wordprocessingml" in content_type:
                try:
                    import io
                    from docx import Document  # type: ignore[import-untyped]

                    doc = Document(io.BytesIO(raw))
                    parts: list[str] = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            parts.append(para.text)
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join(cell.text.strip() for cell in row.cells)
                            if row_text.strip(" |"):
                                parts.append(row_text)
                    text = "\n".join(parts)
                except Exception as exc:
                    return f"Failed to parse DOCX: {exc}"

            elif ext in (".txt", ".csv", ".md") or content_type.startswith("text/"):
                try:
                    text = raw.decode("utf-8", errors="replace")
                except Exception as exc:
                    return f"Failed to decode file as text: {exc}"

            elif ext == ".pdf" or "application/pdf" in content_type:
                size_kb = len(raw) // 1024
                return (
                    f"[PDF file: {filename}, {size_kb} KB] "
                    "PDF text extraction is not available in this service. "
                    "Download the file directly and use a PDF reader."
                )

            else:
                size_kb = len(raw) // 1024
                return f"[Binary file: {filename}, {size_kb} KB — text extraction not supported for this file type]"

            if len(text.encode("utf-8")) > sp_max_content:
                text = text.encode("utf-8")[:sp_max_content].decode("utf-8", errors="ignore")
                text += "\n\n[Content truncated]"

            logger.info("SharePoint downloaded %s — %d chars", filename, len(text))
            return f"# {filename}\n\n{text}" if text else "(empty file)"

        logger.info("Tools MCP: sharepoint tools registered")
